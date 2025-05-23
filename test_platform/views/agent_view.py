import os
import json
import uuid
import pandas as pd
from django.http import JsonResponse, HttpResponse
from rest_framework.views import APIView
from rest_framework.decorators import api_view, permission_classes, authentication_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework_simplejwt.authentication import JWTAuthentication
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from django.conf import settings
from django.db.models import F
from django.utils import timezone
import logging
import docx
import PyPDF2
import io
import time
import requests
from openai import OpenAI
from test_platform.models import APIKey
from django.views.decorators.csrf import csrf_exempt

logger = logging.getLogger(__name__)

# 定义允许的文件类型
ALLOWED_EXTENSIONS = {
    'docx': ['docx'],
    'excel': ['xlsx', 'xls', 'csv'],
    'pdf': ['pdf']
}

# 创建临时文件存储目录
TEMP_DIR = os.path.join(settings.MEDIA_ROOT, 'temp_files')
os.makedirs(TEMP_DIR, exist_ok=True)


class FileParseView(APIView):
    """
    文件解析视图
    处理不同类型文件的上传和解析功能
    """
    permission_classes = [IsAuthenticated]
    authentication_classes = [JWTAuthentication]

    def post(self, request):
        """
        处理文件上传及解析
        :param request: 包含文件和解析参数的请求
        :return: 解析结果或错误信息
        """
        try:
            logger.info("开始处理文件上传请求")
            
            # 检查是否有文件上传
            if 'file' not in request.FILES:
                logger.warning("未找到上传的文件")
                return JsonResponse({
                    'code': 400,
                    'message': '未找到上传的文件',
                    'data': None
                }, status=400)

            uploaded_file = request.FILES['file']
            logger.info(f"接收到文件：{uploaded_file.name}，大小：{uploaded_file.size} 字节")
            
            # 获取文件扩展名并检查是否支持
            file_extension = uploaded_file.name.split('.')[-1].lower()
            logger.debug(f"文件扩展名：{file_extension}")
            
            # 根据文件类型调用相应的解析方法
            if file_extension in ALLOWED_EXTENSIONS['docx']:
                logger.info(f"开始解析Word文档：{uploaded_file.name}")
                result = self.parse_docx(uploaded_file)
                logger.info(f"Word文档解析完成：{uploaded_file.name}")
            elif file_extension in ALLOWED_EXTENSIONS['excel']:
                logger.info(f"开始解析Excel文件：{uploaded_file.name}")
                result = self.parse_excel(uploaded_file, file_extension)
                logger.info(f"Excel文件解析完成：{uploaded_file.name}")
            elif file_extension in ALLOWED_EXTENSIONS['pdf']:
                logger.info(f"开始解析PDF文件：{uploaded_file.name}")
                result = self.parse_pdf(uploaded_file)
                logger.info(f"PDF文件解析完成：{uploaded_file.name}")
            else:
                logger.warning(f"不支持的文件类型: {file_extension}")
                return JsonResponse({
                    'code': 400,
                    'message': f'不支持的文件类型: {file_extension}',
                    'data': None
                }, status=400)
            
            logger.debug(f"文件解析结果：文件类型={result.get('file_type')}, sheets数量={len(result.get('sheets', []))}")
            
            # 检查是否需要发送到DeepSeek处理 - 支持多种数据来源
            send_to_deepseek_param = request.POST.get('send_to_deepseek')
            if send_to_deepseek_param is None and hasattr(request, 'data'):
                send_to_deepseek_param = request.data.get('send_to_deepseek')
            
            # 默认设为true，除非明确指定为false
            send_to_deepseek = False if send_to_deepseek_param is not None and send_to_deepseek_param.lower() == 'false' else True
            logger.debug(f"是否发送到DeepSeek: {send_to_deepseek}")
            
            # 获取API密钥ID - 支持多种数据来源
            api_key_id = request.POST.get('api_key_id')
            if api_key_id is None and hasattr(request, 'data'):
                api_key_id = request.data.get('api_key_id')
            
            deepseek_response = None
            if send_to_deepseek:
                logger.info("开始将解析结果发送到DeepSeek进行分析")
                deepseek_response = self.send_to_deepseek(result, request.user, api_key_id)
                # 添加DeepSeek响应到结果中
                if deepseek_response:
                    logger.info("成功接收DeepSeek分析结果")
                    logger.debug(f"DeepSeek状态: {deepseek_response.get('status')}, 令牌数: {deepseek_response.get('total_tokens', 0)}")
                    result['deepseek_response'] = deepseek_response
                    
                    # 如果deepseek_response中包含test_cases，将其提升到顶层
                    if 'test_cases' in deepseek_response:
                        logger.info(f"从DeepSeek响应中提取到{len(deepseek_response['test_cases'])}个测试用例")
                        result['test_cases'] = deepseek_response['test_cases']
                    
                    # 如果不存在test_cases字段，则尝试从内容中解析
                    elif 'content' in deepseek_response:
                        logger.info("尝试从DeepSeek返回内容中解析测试用例")
                        # 尝试从内容中解析测试用例
                        test_cases = parse_test_cases_from_text(deepseek_response['content'], 1)  # 默认使用项目ID 1
                        if test_cases:
                            logger.info(f"成功从内容中解析出{len(test_cases)}个测试用例")
                            result['test_cases'] = test_cases
                            # 也添加到deepseek_response中保持一致性
                            deepseek_response['test_cases'] = test_cases
                        else:
                            logger.warning("无法从DeepSeek内容中解析出测试用例")
                else:
                    logger.warning("发送到DeepSeek失败或未返回有效结果")
            
            # 确保所有文件类型都至少有一个空的test_cases字段
            if 'test_cases' not in result:
                logger.debug("添加默认空的test_cases字段")
                result['test_cases'] = []
            
            logger.info("文件解析请求处理完成")
            return JsonResponse({
                'code': 200,
                'message': '文件解析成功',
                'data': result
            })
            
        except Exception as e:
            logger.error(f"文件解析失败: {str(e)}", exc_info=True)
            return JsonResponse({
                'code': 500,
                'message': f'文件解析失败: {str(e)}',
                'data': None
            }, status=500)
    
    def send_to_deepseek(self, parsed_data, user, api_key_id=None):
        """
        将解析后的数据发送至DeepSeek API
        :param parsed_data: 解析后的文件数据
        :param user: 当前用户
        :param api_key_id: 可选，指定使用的API密钥ID
        :return: DeepSeek API的响应结果
        """
        try:
            logger.info(f"开始准备发送数据到DeepSeek，用户ID: {user.id}")
            
            # 获取API密钥
            api_key_obj = None
            
            # 如果提供了api_key_id，直接使用该API密钥
            if api_key_id:
                logger.debug(f"使用指定的API密钥ID: {api_key_id}")
                try:
                    api_key_obj = APIKey.objects.get(api_key_id=api_key_id, user=user, is_active=True,
                                                    service_type='deepseek')
                    logger.debug(f"成功获取指定的API密钥: {api_key_obj.key_name}")
                except APIKey.DoesNotExist:
                    logger.warning(f"指定的DeepSeek API密钥不存在: {api_key_id}")
                    return None
            else:
                # 否则，查找该用户的默认DeepSeek API密钥
                logger.debug("未指定API密钥ID，尝试使用默认API密钥")
                api_key_obj = APIKey.objects.filter(
                    user=user,
                    service_type='deepseek',
                    is_active=True,
                    is_default=True
                ).first()
                
                # 如果没有默认API密钥，则获取最近创建的DeepSeek API密钥
                if not api_key_obj:
                    logger.debug("未找到默认API密钥，尝试使用最近创建的API密钥")
                    api_key_obj = APIKey.objects.filter(
                        user=user,
                        service_type='deepseek',
                        is_active=True
                    ).order_by('-create_time').first()
            
            # 如果未找到有效的API密钥，返回错误
            if not api_key_obj:
                logger.warning("未找到有效的DeepSeek API密钥")
                return {
                    "error": "未找到有效的DeepSeek API密钥，请先创建DeepSeek API密钥",
                    "status": "failed"
                }
            
            logger.info(f"使用API密钥: {api_key_obj.key_name}")
            
            # 获取额外配置参数
            config = {}
            if hasattr(api_key_obj, 'config') and api_key_obj.config:
                try:
                    config = json.loads(api_key_obj.config)
                    logger.debug(f"读取API密钥配置: 模型={config.get('model', 'deepseek-chat')}, 温度={config.get('temperature', 0.7)}")
                except json.JSONDecodeError:
                    logger.warning(f"API密钥配置JSON解析失败: {api_key_obj.config}")
            
            # 提取配置参数
            model = config.get('model', 'deepseek-chat')
            temperature = config.get('temperature', 0.7)
            
            # 准备请求数据
            # 根据文件类型构造不同的消息内容
            file_type = parsed_data.get('file_type')
            file_name = parsed_data.get('file_name')
            
            logger.info(f"构建DeepSeek请求消息，文件类型: {file_type}, 文件名: {file_name}")
            
            # 构建发送到DeepSeek的消息
            message_content = f"这是一个解析自{file_name}的{file_type}文件内容:\n\n"
            
            if file_type == 'docx':
                # 添加文档段落内容
                paragraphs = parsed_data.get('content', {}).get('paragraphs', [])
                logger.debug(f"处理Word文档段落: {len(paragraphs)}个段落")
                for para in paragraphs:
                    message_content += f"{para['text']}\n"
                
                # 添加表格内容
                tables = parsed_data.get('content', {}).get('tables', [])
                if tables:
                    logger.debug(f"处理Word文档表格: {len(tables)}个表格")
                    message_content += "\n表格内容：\n"
                    for table in tables:
                        message_content += f"表格 {table['id']}:\n"
                        for row in table['data']:
                            message_content += "|".join(row) + "\n"
                        message_content += "\n"
            
            elif file_type in ['xlsx', 'xls', 'csv']:
                # 添加工作表内容
                sheets = parsed_data.get('sheets', [])
                logger.debug(f"处理Excel工作表: {len(sheets)}个工作表")
                for sheet in sheets:
                    message_content += f"\n工作表: {sheet['name']}\n"
                    
                    # 如果有数据
                    if sheet['data']:
                        # 获取所有列名
                        columns = sheet['data'][0].keys()
                        
                        # 添加列标题
                        message_content += "|".join(columns) + "\n"
                        
                        # 添加分隔行
                        message_content += "|".join(["-" * len(col) for col in columns]) + "\n"
                        
                        # 添加数据行
                        for row in sheet['data']:
                            message_content += "|".join([str(row.get(col, "")) for col in columns]) + "\n"
                    
                    message_content += "\n"
            
            elif file_type == 'pdf':
                # 添加PDF页面内容
                pages = parsed_data.get('content', {}).get('pages', [])
                logger.debug(f"处理PDF页面: {len(pages)}个页面")
                for page in pages:
                    message_content += f"\n页码 {page['page_num']}:\n"
                    message_content += f"{page['text']}\n"
            
            logger.debug(f"消息内容构建完成，长度: {len(message_content)} 字符")
            
            # 使用OpenAI库调用DeepSeek API
            logger.info(f"初始化OpenAI客户端，连接到: {api_key_obj.api_base}")
            client = OpenAI(
                api_key=api_key_obj.api_key,
                base_url=api_key_obj.api_base  # DeepSeek 使用 https://api.deepseek.com
            )
            
            # 调用DeepSeek API
            logger.info(f"开始调用DeepSeek API，模型: {model}, 温度: {temperature}")
            start_time = time.time()
            
            # 构建系统提示信息
            system_prompt = """
# 接口测试用例生成指南

你是一个专业的API测试工程师，需要从输入的文件内容中提取测试用例信息，并生成标准格式的接口测试用例。

## 输出格式要求

必须且只能以下面的JSON格式输出测试用例，这是系统能识别的唯一格式：

```json
[
  {
    "title": "测试用例名称", 
    "api_path": "接口路径，如http://localhost:8081/api/login/",
    "method": "请求方式，如GET、POST、PUT、DELETE",
    "priority": "优先级，如高、中、低",
    "headers": {
      "content-type": "application/json"
    },
    "body": {
      "key1": "value1",
      "key2": "value2"
    },
    "assertions": "断言类型",
    "expected_result": {
      "code": 200,
      "message": "预期响应消息"
    },
    "extractors": [
      {
        "name": "token",
        "type": "jsonpath",
        "expression": "$.data.token",
        "defaultValue": "提取失败",
        "enabled": true
      }
    ],
    "tests": []
  }
]
```

## 注意事项

1. 识别文件中的测试用例关键信息，包括：用例名称、接口路径、请求方法、优先级、请求体、预期结果等
2. 优先级映射：P0/P1=>高，P2=>中，P3=>低
3. 请求体和预期结果必须是有效的JSON格式
4. 不要添加任何无关解释或说明，只输出JSON格式的测试用例
5. 提取的测试用例应放在一个数组中返回，即使只有一个用例
6. 确保生成的测试用例符合上述格式规范

请分析以下文件内容，生成符合要求的测试用例：
"""
            
            # 使用分批处理函数处理大型数据
            content = process_large_data_with_llm(
                message_content,
                client,
                model,
                temperature,
                max_chunk_size=8000,  # 根据模型token限制调整
                system_prompt=system_prompt
            )
            
            end_time = time.time()
            logger.info(f"DeepSeek API调用完成，耗时: {end_time - start_time:.2f} 秒")
            
            # 更新API密钥使用统计
            logger.debug(f"更新API密钥使用统计，密钥ID: {api_key_obj.api_key_id}")
            APIKey.objects.filter(api_key_id=api_key_obj.api_key_id).update(
                usage_count=F('usage_count') + 1,
                last_used_at=timezone.now()
            )
            
            # 构造响应对象
            response_id = str(uuid.uuid4())
            deepseek_response = {
                "status": "success",
                "model": model,
                "content": content,  # 使用处理后的内容
                "id": response_id,
                "created": int(time.time()),
                "finish_reason": "stop",
                "prompt_tokens": len(message_content) // 4,  # 估算token数量
                "completion_tokens": len(content) // 4,  # 估算token数量
                "total_tokens": (len(message_content) + len(content)) // 4  # 估算token数量
            }
            
            logger.info(f"DeepSeek响应处理完成，总令牌数: {deepseek_response['total_tokens']}")
            
            # 生成测试用例
            logger.info("开始生成测试用例")
            test_cases = self.generate_test_cases(deepseek_response, parsed_data)
            if test_cases:
                logger.info(f"测试用例生成成功，数量: {len(test_cases)}")
                deepseek_response["test_cases"] = test_cases
            else:
                logger.warning("未能生成任何测试用例")
            
            return deepseek_response
            
        except Exception as e:
            logger.error(f"发送数据到DeepSeek失败: {str(e)}", exc_info=True)
            return {
                "error": str(e),
                "status": "failed"
            }
    
    def generate_test_cases(self, deepseek_response, parsed_data):
        """
        根据DeepSeek的分析结果生成测试用例
        :param deepseek_response: DeepSeek的分析结果
        :param parsed_data: 解析后的文件数据
        :return: 测试用例列表
        """
        try:
            logger.info("开始从DeepSeek响应生成测试用例")
            
            # 提取文件类型和内容
            file_type = parsed_data.get('file_type')
            logger.debug(f"文件类型: {file_type}")
            test_cases = []
            
            # 根据文件类型处理
            if file_type in ['xlsx', 'xls', 'csv']:
                sheets = parsed_data.get('sheets', [])
                logger.debug(f"处理{len(sheets)}个工作表的数据")
                
                # 遍历所有工作表
                for sheet in sheets:
                    sheet_name = sheet.get('name', '')
                    rows = sheet.get('data', [])
                    
                    # 跳过空工作表
                    if not rows:
                        logger.debug(f"工作表 '{sheet_name}' 为空，跳过")
                        continue
                    
                    logger.debug(f"处理工作表 '{sheet_name}' 中的{len(rows)}行数据")
                    
                    # 遍历每行数据，构建测试用例
                    for i, row in enumerate(rows):
                        # 构建基本测试用例模板
                        case_title = row.get('用例名称', '') or row.get('case_name', '') or ''
                        if not case_title:
                            logger.debug(f"跳过第{i+1}行，未找到有效的用例名称")
                            continue
                            
                        logger.debug(f"处理测试用例: {case_title}")
                        
                        test_case = {
                            "title": case_title,
                            "api_path": row.get('用例路径', '') or row.get('case_path', '') or '',
                            "method": row.get('请求方法', '') or row.get('case_request_method', '') or 'GET',
                            "priority": self._map_priority(row.get('优先级', '') or row.get('case_priority', '') or '中'),
                            "headers": {
                                "content-type": "application/json"
                            },
                            "body": self._parse_body(row.get('请求体', '') or row.get('case_requests_body', '') or '{}'),
                            "assertions": row.get('断言类型', '') or row.get('case_assert_type', '') or '',
                            "expected_result": self._parse_expected_result(row.get('预期结果', '') or row.get('case_expect_result', '') or '{}'),
                            "project_id": "1",  # 默认项目ID
                            "body_type": "raw",
                            "raw_content_type": "application/json",
                            "form_data": [],
                            "extractors": self._generate_extractors(row),
                            "tests": []
                        }
                        
                        # 如果有描述，添加描述
                        description = row.get('用例描述', '') or row.get('case_description', '')
                        if description:
                            test_case["description"] = description
                        
                        # 添加到测试用例列表
                        test_cases.append(test_case)
                        logger.debug(f"成功创建测试用例: {test_case['title']}, 方法: {test_case['method']}, 路径: {test_case['api_path']}")
            
            logger.info(f"测试用例生成完成，共生成{len(test_cases)}个测试用例")
            return test_cases
            
        except Exception as e:
            logger.error(f"生成测试用例失败: {str(e)}", exc_info=True)
            return []
    
    def _map_priority(self, priority):
        """
        映射优先级
        :param priority: 原始优先级值
        :return: 映射后的优先级值
        """
        priority_map = {
            'P0': '最高',
            'P1': '高',
            'P2': '中',
            'P3': '低',
            '高': '高',
            '中': '中',
            '低': '低'
        }
        return priority_map.get(priority, '中')
    
    def _parse_body(self, body_str):
        """
        解析请求体
        :param body_str: 请求体字符串
        :return: 请求体对象
        """
        if not body_str:
            return {}
            
        try:
            # 尝试解析为JSON
            return json.loads(body_str)
        except json.JSONDecodeError:
            # 如果解析失败，尝试将字符串解析为字典
            body = {}
            # 简单解析格式如 "key1=value1\nkey2=value2"
            for line in body_str.split('\n'):
                if '=' in line:
                    key, value = line.split('=', 1)
                    body[key.strip()] = value.strip()
            return body
    
    def _parse_expected_result(self, result_str):
        """
        解析预期结果
        :param result_str: 预期结果字符串
        :return: 预期结果对象
        """
        if not result_str:
            return {}
            
        try:
            # 尝试解析为JSON
            return json.loads(result_str)
        except json.JSONDecodeError:
            # 如果解析失败，尝试将字符串解析为字典
            result = {}
            # 简单解析格式如 "key1: value1\nkey2: value2"
            for line in result_str.split('\n'):
                if ':' in line:
                    key, value = line.split(':', 1)
                    result[key.strip()] = value.strip()
            return result
    
    def _generate_extractors(self, row):
        """
        生成提取器列表
        :param row: 数据行
        :return: 提取器列表
        """
        # 默认返回一个空的提取器列表
        extractors = []
        
        # 如果有设置提取器字段，解析提取器
        extractors_str = row.get('提取器', '') or row.get('case_extractors', '')
        if extractors_str:
            try:
                # 尝试解析为JSON
                parsed_extractors = json.loads(extractors_str)
                if isinstance(parsed_extractors, list):
                    return parsed_extractors
                elif isinstance(parsed_extractors, dict):
                    extractors.append(parsed_extractors)
            except json.JSONDecodeError:
                # 如果解析失败，尝试简单解析
                # 简单提取器格式例如: "name=token,type=jsonpath,expression=$.data.token"
                for extractor_str in extractors_str.split(';'):
                    extractor = {}
                    for part in extractor_str.split(','):
                        if '=' in part:
                            key, value = part.split('=', 1)
                            extractor[key.strip()] = value.strip()
                    if extractor and 'name' in extractor:
                        extractors.append({
                            "name": extractor.get('name', ''),
                            "type": extractor.get('type', 'jsonpath'),
                            "expression": extractor.get('expression', ''),
                            "defaultValue": extractor.get('defaultValue', ''),
                            "enabled": True
                        })
        
        # 如果没有提取器，添加默认提取器
        if not extractors:
            extractors = [
                {
                    "name": "token",
                    "type": "jsonpath",
                    "expression": "$.data.token",
                    "defaultValue": "提取失败",
                    "enabled": True
                }
            ]
            
        return extractors
    
    def save_temp_file(self, file):
        """
        保存临时文件
        :param file: 上传的文件对象
        :return: 临时文件的路径
        """
        # 生成唯一的文件名
        filename = f"{uuid.uuid4()}_{file.name}"
        filepath = os.path.join(TEMP_DIR, filename)
        
        # 保存文件
        with open(filepath, 'wb+') as destination:
            for chunk in file.chunks():
                destination.write(chunk)
                
        return filepath
    
    def parse_docx(self, file):
        """
        解析Word文档
        :param file: 上传的docx文件
        :return: 解析结果，包含文档内容和结构信息
        """
        logger.info(f"开始解析Word文档: {file.name}")
        
        # 使用python-docx库解析文档
        try:
            doc = docx.Document(file)
            logger.debug("成功加载Word文档")
            
            # 提取文档内容
            content = {
                'paragraphs': [],
                'tables': []
            }
            
            # 提取段落
            logger.debug(f"开始提取文档段落，共{len(doc.paragraphs)}个段落")
            for para in doc.paragraphs:
                if para.text.strip():
                    content['paragraphs'].append({
                        'text': para.text,
                        'style': para.style.name if para.style else 'Normal'
                    })
            logger.debug(f"提取了{len(content['paragraphs'])}个非空段落")
            
            # 提取表格
            logger.debug(f"开始提取文档表格，共{len(doc.tables)}个表格")
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text)
                    table_data.append(row_data)
                content['tables'].append({
                    'id': i + 1,
                    'data': table_data
                })
            logger.debug(f"提取了{len(content['tables'])}个表格")
            
            # 将段落和表格内容转换为表格式数据
            sheets_data = []
            
            # 将段落转换为表格数据
            logger.debug("开始将段落转换为表格数据")
            paragraphs_data = []
            for i, para in enumerate(content['paragraphs']):
                paragraphs_data.append({
                    'ID': i + 1,
                    '段落内容': para['text'],
                    '样式': para['style']
                })
            
            # 将表格内容转换为统一格式
            logger.debug("开始将表格转换为统一格式")
            tables_data = []
            for table in content['tables']:
                if table['data'] and len(table['data']) > 0:
                    # 获取表头
                    headers = ["列" + str(i+1) for i in range(len(table['data'][0]))]
                    
                    # 构建表格数据
                    for row_idx, row in enumerate(table['data']):
                        row_dict = {'行号': row_idx + 1}
                        for col_idx, cell in enumerate(row):
                            row_dict[headers[col_idx]] = cell
                        tables_data.append(row_dict)
            
            # 将所有内容整合到sheets中
            sheets = [
                {
                    'name': '文档段落',
                    'data': paragraphs_data
                }
            ]
            
            # 如果有表格数据，添加到sheets中
            if tables_data:
                sheets.append({
                    'name': '文档表格',
                    'data': tables_data
                })
            
            logger.info(f"Word文档解析完成: {file.name}，生成了{len(sheets)}个工作表")
            
            # 返回解析结果
            return {
                'file_type': 'docx',
                'file_name': file.name,
                'content': content,  # 保留原始content字段以保持兼容性
                'sheets': sheets     # 添加sheets字段以统一格式
            }
        except Exception as e:
            logger.error(f"解析Word文档失败: {str(e)}", exc_info=True)
            raise
    
    def parse_excel(self, file, file_extension):
        """
        解析Excel文件
        :param file: 上传的Excel文件
        :param file_extension: 文件扩展名
        :return: 解析结果，包含工作表和数据
        """
        logger.info(f"开始解析Excel文件: {file.name}, 扩展名: {file_extension}")
        
        # 保存临时文件
        temp_file_path = self.save_temp_file(file)
        logger.debug(f"创建临时文件: {temp_file_path}")
        
        excel_file = None
        df = None
        
        try:
            # 根据文件类型选择读取方式
            if file_extension == 'csv':
                logger.debug("使用CSV读取方式")
                # 读取CSV文件
                df = pd.read_csv(temp_file_path)
                logger.debug(f"成功读取CSV文件，包含{len(df)}行数据")
                # 确保DataFrame资源被释放
                data = df.fillna('').to_dict(orient='records')
                df = None  # 显式释放DataFrame
                
                sheets = [{
                    'name': 'Sheet1',
                    'data': data
                }]
                logger.debug(f"CSV数据转换为单个工作表，包含{len(data)}行数据")
            else:
                logger.debug("使用Excel读取方式")
                # 读取Excel文件
                excel_file = pd.ExcelFile(temp_file_path)
                sheets = []
                
                logger.debug(f"Excel文件包含{len(excel_file.sheet_names)}个工作表")
                
                # 遍历所有工作表
                for sheet_name in excel_file.sheet_names:
                    logger.debug(f"读取工作表: {sheet_name}")
                    df = excel_file.parse(sheet_name)
                    logger.debug(f"工作表 {sheet_name} 包含{len(df)}行数据")
                    data = df.fillna('').to_dict(orient='records')
                    df = None  # 显式释放DataFrame
                    sheets.append({
                        'name': sheet_name,
                        'data': data
                    })
                
                # 关闭Excel文件
                if excel_file is not None:
                    excel_file.close()
                    excel_file = None
                    logger.debug("已关闭Excel文件")
            
            logger.info(f"Excel文件解析完成: {file.name}，生成了{len(sheets)}个工作表")
            
            # 返回解析结果
            return {
                'file_type': file_extension,
                'file_name': file.name,
                'sheets': sheets
            }
            
        finally:
            # 确保资源被释放
            if df is not None:
                del df
            if excel_file is not None:
                excel_file.close()
            
            # 添加短暂延迟确保文件解锁
            time.sleep(0.5)
            
            # 删除临时文件
            try:
                if os.path.exists(temp_file_path):
                    os.remove(temp_file_path)
                    logger.debug(f"已删除临时文件: {temp_file_path}")
            except Exception as e:
                logger.warning(f"无法删除临时文件 {temp_file_path}: {str(e)}")
                # 继续执行，不影响主要功能
    
    def parse_pdf(self, file):
        """
        解析PDF文件
        :param file: 上传的PDF文件
        :return: 解析结果，包含文档页面和内容
        """
        logger.info(f"开始解析PDF文件: {file.name}")
        
        try:
            # 从上传的文件创建二进制流
            file_stream = io.BytesIO(file.read())
            
            # 使用PyPDF2解析PDF
            pdf_reader = PyPDF2.PdfReader(file_stream)
            num_pages = len(pdf_reader.pages)
            logger.debug(f"PDF文件包含{num_pages}页")
            
            # 提取文档内容
            content = {
                'num_pages': num_pages,
                'pages': []
            }
            
            # 提取每页内容
            for i in range(num_pages):
                logger.debug(f"解析第{i+1}页")
                page = pdf_reader.pages[i]
                page_text = page.extract_text()
                content['pages'].append({
                    'page_num': i + 1,
                    'text': page_text
                })
            
            # 将PDF内容转换为表格式数据
            logger.debug("将PDF内容转换为表格数据")
            pdf_data = []
            for page in content['pages']:
                pdf_data.append({
                    '页码': page['page_num'],
                    '内容': page['text']
                })
            
            # 构建sheets结构
            sheets = [
                {
                    'name': 'PDF内容',
                    'data': pdf_data
                }
            ]
            
            logger.info(f"PDF文件解析完成: {file.name}")
            
            # 返回解析结果
            return {
                'file_type': 'pdf',
                'file_name': file.name,
                'content': content,  # 保留原始content字段以保持兼容性
                'sheets': sheets     # 添加sheets字段以统一格式
            }
        except Exception as e:
            logger.error(f"解析PDF文件失败: {str(e)}", exc_info=True)
            raise


@api_view(['POST'])
@permission_classes([IsAuthenticated])
@authentication_classes([JWTAuthentication])
def file_parse(request):
    """
    文件解析API视图函数
    :param request: HTTP请求对象
    :return: JsonResponse包含解析结果
    """
    # 实例化视图类并调用post方法
    view = FileParseView()
    return view.post(request)

@csrf_exempt
@api_view(['POST', 'GET'])
@permission_classes([IsAuthenticated])
@authentication_classes([JWTAuthentication])
def save_deepseek_config(request):
    """
    保存DeepSeek API配置
    接收API密钥和模型参数，保存到用户的API密钥配置中
    
    请求参数:
    - apiKey: DeepSeek API密钥
    - model: 使用的模型名称，例如 'deepseek-chat'
    - temperature: 温度参数，控制生成内容的随机性
    
    返回:
    - 保存成功或失败的状态信息
    """
    # 处理GET请求
    if request.method == 'GET':
        return get_deepseek_config(request)
        
    # 处理POST请求
    try:
        # 解析请求体
        data = json.loads(request.body)
        api_key = data.get('apiKey')
        model = data.get('model', 'deepseek-chat')
        temperature = data.get('temperature', 0.7)
        
        # 验证必要参数
        if not api_key:
            return JsonResponse({
                'code': 400,
                'message': 'API密钥不能为空',
                'data': None
            }, status=400)
            
        # 检查模型是否有效
        valid_models = ['deepseek-chat', 'deepseek-reasoner']
        if model not in valid_models:
            return JsonResponse({
                'code': 400,
                'message': f'不支持的模型名称。支持的模型: {", ".join(valid_models)}',
                'data': None
            }, status=400)
        
        # 检查温度参数是否有效
        try:
            temperature = float(temperature)
            if not (0 <= temperature <= 1):
                return JsonResponse({
                    'code': 400,
                    'message': '温度参数应该在0到1之间',
                    'data': None
                }, status=400)
        except (ValueError, TypeError):
            return JsonResponse({
                'code': 400,
                'message': '温度参数应该是一个浮点数',
                'data': None
            }, status=400)
            
        # 生成配置名称
        key_name = f"deepseek-{model}-{timezone.now().strftime('%Y%m%d%H%M%S')}"
        
        # 检查用户是否已有DeepSeek配置
        has_default = APIKey.objects.filter(
            user=request.user,
            service_type='deepseek',
            is_default=True
        ).exists()
        
        # 创建或更新API密钥记录
        api_key_obj, created = APIKey.objects.update_or_create(
            user=request.user,
            service_type='deepseek',
            is_default=True,
            defaults={
                'key_name': key_name,
                'api_key': api_key,
                'api_base': 'https://api.deepseek.com',
                'is_active': True,
                'usage_count': 0,
                'last_used_at': None,
            }
        )
        
        # 保存额外参数到配置中
        api_key_obj.config = json.dumps({
            'model': model,
            'temperature': temperature
        })
        api_key_obj.save()
        
        # 返回成功响应
        return JsonResponse({
            'code': 200,
            'message': '已成功' + ('创建' if created else '更新') + ' DeepSeek API配置',
            'data': {
                'api_key_id': api_key_obj.api_key_id,
                'key_name': api_key_obj.key_name,
                'service_type': 'deepseek',
                'model': model,
                'temperature': temperature,
                'is_default': api_key_obj.is_default,
                'create_time': api_key_obj.create_time.strftime('%Y-%m-%d %H:%M:%S')
            }
        })
        
    except Exception as e:
        logger.error(f"保存DeepSeek配置失败: {str(e)}", exc_info=True)
        return JsonResponse({
            'code': 500,
            'message': f'保存DeepSeek配置失败: {str(e)}',
            'data': None
        }, status=500)

@api_view(['GET'])
@permission_classes([IsAuthenticated])
@authentication_classes([JWTAuthentication])
def get_deepseek_config(request):
    """
    获取当前用户的DeepSeek API配置
    
    返回:
    - 当前用户的DeepSeek配置信息
    """
    try:
        # 获取用户的DeepSeek API密钥配置
        api_key_obj = APIKey.objects.filter(
            user=request.user,
            service_type='deepseek',
            is_active=True,
            is_default=True
        ).first()
        
        # 如果没找到默认配置，尝试获取最近创建的配置
        if not api_key_obj:
            api_key_obj = APIKey.objects.filter(
                user=request.user,
                service_type='deepseek',
                is_active=True
            ).order_by('-create_time').first()
        
        # 如果找不到配置，返回空数据
        if not api_key_obj:
            return JsonResponse({
                'code': 200,
                'message': '未找到DeepSeek API配置',
                'data': {
                    'has_config': False
                }
            })
            
        # 解析额外配置
        config = {}
        if hasattr(api_key_obj, 'config') and api_key_obj.config:
            try:
                config = json.loads(api_key_obj.config)
            except json.JSONDecodeError:
                pass
        
        # 返回配置信息
        return JsonResponse({
            'code': 200,
            'message': '获取DeepSeek配置成功',
            'data': {
                'has_config': True,
                'api_key_id': api_key_obj.api_key_id,
                'key_name': api_key_obj.key_name,
                'api_base': api_key_obj.api_base,
                'model': config.get('model', 'deepseek-chat'),
                'temperature': config.get('temperature', 0.7),
                'create_time': api_key_obj.create_time.strftime('%Y-%m-%d %H:%M:%S'),
                'last_used_at': api_key_obj.last_used_at.strftime('%Y-%m-%d %H:%M:%S') if api_key_obj.last_used_at else None,
                'usage_count': api_key_obj.usage_count
            }
        })
    
    except Exception as e:
        logger.error(f"获取DeepSeek配置失败: {str(e)}", exc_info=True)
        return JsonResponse({
            'code': 500,
            'message': f'获取DeepSeek配置失败: {str(e)}',
            'data': None
        }, status=500)

@api_view(['POST'])
@permission_classes([IsAuthenticated])
@authentication_classes([JWTAuthentication])
def create_test_case(request):
    """
    根据DeepSeek分析结果创建测试用例
    
    请求参数:
    - deepseek_response: DeepSeek分析结果
    - project_id: 项目ID
    
    返回:
    - 创建的测试用例信息
    """
    try:
        logger.info(f"接收到创建测试用例请求，用户ID: {request.user.id}")
        
        # 解析请求数据
        deepseek_response = request.data.get('deepseek_response')
        project_id = request.data.get('project_id', 1)
        logger.debug(f"项目ID: {project_id}")
        
        if not deepseek_response:
            logger.warning("请求中未提供DeepSeek分析结果")
            return JsonResponse({
                'code': 400,
                'message': 'DeepSeek分析结果不能为空',
                'data': None
            }, status=400)
        
        # 导入TestCase模型
        from test_platform.models import TestCase
        
        # 如果deepseek_response中已有test_cases字段，直接使用
        test_cases = deepseek_response.get('test_cases', [])
        logger.debug(f"从DeepSeek响应中直接获取到{len(test_cases)}个测试用例")
        
        # 如果没有test_cases字段，尝试从文本内容解析测试用例
        if not test_cases and 'content' in deepseek_response:
            logger.info("尝试从DeepSeek内容中解析测试用例")
            # 尝试从内容中解析测试用例
            test_cases = parse_test_cases_from_text(deepseek_response['content'], project_id)
            logger.debug(f"从内容中解析出{len(test_cases)}个测试用例")
        
        # 如果仍然没有测试用例，返回错误
        if not test_cases:
            logger.warning("无法从DeepSeek分析结果中解析测试用例")
            return JsonResponse({
                'code': 400,
                'message': '无法从DeepSeek分析结果中解析测试用例',
                'data': None
            }, status=400)
        
        # 创建测试用例
        created_cases = []
        logger.info(f"开始创建{len(test_cases)}个测试用例")
        
        for i, case in enumerate(test_cases):
            # 确保case_name不为空
            if not case.get('title'):
                logger.warning(f"跳过第{i+1}个测试用例，标题为空")
                continue
                
            # 添加project_id
            case['project_id'] = project_id
            
            logger.debug(f"创建测试用例: {case.get('title')}")
            
            # 创建测试用例
            try:
                test_case = TestCase.objects.create(
                    case_name=case.get('title', ''),
                    case_description=case.get('description', ''),
                    case_path=case.get('api_path', ''),
                    case_request_method=case.get('method', 'GET'),
                    case_priority=case.get('priority', '中'),
                    case_params='',
                    case_status=1,  # 假设1表示"启用"
                    case_precondition='',
                    case_request_headers=json.dumps(case.get('headers', {})),
                    case_requests_body=json.dumps(case.get('body', {})),
                    case_assert_type=case.get('assertions', ''),
                    case_assert_contents=json.dumps(case.get('tests', [])),
                    case_expect_result=json.dumps(case.get('expected_result', {})),
                    create_time=timezone.now(),
                    update_time=timezone.now(),
                    creator_id=request.user.id,
                    project_id=project_id,
                    case_extractors=json.dumps(case.get('extractors', []))
                )
                
                created_cases.append({
                    'case_id': test_case.id,
                    'case_name': test_case.case_name
                })
                logger.debug(f"测试用例创建成功: ID={test_case.id}, 名称={test_case.case_name}")
            except Exception as e:
                logger.error(f"创建测试用例失败: {str(e)}")
        
        logger.info(f"测试用例创建完成，成功创建{len(created_cases)}个测试用例")
        return JsonResponse({
            'code': 200,
            'message': f'成功创建{len(created_cases)}个测试用例',
            'data': {
                'cases': created_cases
            }
        })
        
    except Exception as e:
        logger.error(f"创建测试用例失败: {str(e)}", exc_info=True)
        return JsonResponse({
            'code': 500,
            'message': f'创建测试用例失败: {str(e)}',
            'data': None
        }, status=500)


def parse_test_cases_from_text(content, project_id):
    """
    从文本内容中解析测试用例
    :param content: DeepSeek分析结果文本内容
    :param project_id: 项目ID
    :return: 测试用例列表
    """
    test_cases = []
    
    try:
        logger.info(f"开始从文本内容解析测试用例，项目ID: {project_id}")
        logger.debug(f"内容长度: {len(content)} 字符")
        
        # 首先尝试直接解析整个内容为JSON
        try:
            logger.debug("尝试将整个内容解析为JSON")
            json_content = json.loads(content)
            if isinstance(json_content, list):
                # 如果是测试用例数组，直接添加
                logger.info(f"成功解析为JSON数组，包含{len(json_content)}个测试用例")
                for case in json_content:
                    if isinstance(case, dict):
                        case['project_id'] = project_id
                        test_cases.append(case)
                return test_cases
            elif isinstance(json_content, dict):
                # 如果是单个测试用例，添加到列表
                logger.info("成功解析为单个JSON测试用例")
                json_content['project_id'] = project_id
                test_cases.append(json_content)
                return test_cases
        except json.JSONDecodeError:
            # 如果整体不是JSON，继续下面的解析
            logger.debug("整个内容不是有效的JSON，尝试其他解析方法")
            pass
        
        # 尝试寻找JSON格式的测试用例
        import re
        
        # 匹配包含在```json和```之间的JSON内容
        logger.debug("尝试提取Markdown代码块中的JSON内容")
        json_pattern = r'```json\s*([\s\S]*?)\s*```'
        matches = re.findall(json_pattern, content)
        
        if matches:
            logger.debug(f"找到{len(matches)}个JSON代码块")
        
        for i, match in enumerate(matches):
            try:
                logger.debug(f"解析第{i+1}个JSON代码块")
                case_data = json.loads(match)
                if isinstance(case_data, dict):
                    # 单个测试用例
                    logger.debug("解析为单个测试用例")
                    case_data['project_id'] = project_id
                    test_cases.append(case_data)
                elif isinstance(case_data, list):
                    # 测试用例列表
                    logger.debug(f"解析为测试用例列表，包含{len(case_data)}个测试用例")
                    for case in case_data:
                        if isinstance(case, dict):
                            case['project_id'] = project_id
                            test_cases.append(case)
            except json.JSONDecodeError:
                logger.warning(f"第{i+1}个代码块解析失败，不是有效的JSON")
                continue
        
        # 如果没有找到JSON格式的测试用例，尝试寻找Markdown表格格式的测试用例
        if not test_cases:
            logger.debug("尝试从Markdown格式中解析测试用例")
            # 将内容分割成行
            lines = content.split('\n')
            current_case = None
            
            for i, line in enumerate(lines):
                # 尝试识别标题行
                if line.startswith('## ') or line.startswith('### '):
                    # 保存上一个用例（如果有）
                    if current_case and 'title' in current_case:
                        test_cases.append(current_case)
                        logger.debug(f"保存测试用例: {current_case['title']}")
                    
                    title = line.strip('#').strip()
                    logger.debug(f"发现新测试用例标题: {title}")
                    
                    # 创建新用例
                    current_case = {
                        'title': title,
                        'project_id': project_id,
                        'method': 'GET',
                        'priority': '中',
                        'headers': {'content-type': 'application/json'},
                        'body': {},
                        'assertions': '',
                        'expected_result': {},
                        'body_type': 'raw',
                        'raw_content_type': 'application/json',
                        'form_data': [],
                        'extractors': [],
                        'tests': []
                    }
                elif current_case and ':' in line:
                    # 解析键值对
                    key, value = line.split(':', 1)
                    key = key.strip().lower()
                    value = value.strip()
                    
                    logger.debug(f"解析属性: {key}={value}")
                    
                    if key == 'api_path' or key == '接口路径':
                        current_case['api_path'] = value
                    elif key == 'method' or key == '请求方法':
                        current_case['method'] = value
                    elif key == 'priority' or key == '优先级':
                        current_case['priority'] = value
                    elif key == 'body' or key == '请求体' or key == 'request body':
                        try:
                            current_case['body'] = json.loads(value)
                            logger.debug("成功解析请求体为JSON")
                        except json.JSONDecodeError:
                            current_case['body'] = value
                            logger.debug("请求体不是有效的JSON，保存为字符串")
                    elif key == 'expected_result' or key == '预期结果':
                        try:
                            current_case['expected_result'] = json.loads(value)
                            logger.debug("成功解析预期结果为JSON")
                        except json.JSONDecodeError:
                            current_case['expected_result'] = value
                            logger.debug("预期结果不是有效的JSON，保存为字符串")
            
            # 添加最后一个用例
            if current_case and 'title' in current_case:
                test_cases.append(current_case)
                logger.debug(f"保存最后一个测试用例: {current_case['title']}")
        
        logger.info(f"测试用例解析完成，共解析出{len(test_cases)}个测试用例")
    
    except Exception as e:
        logger.error(f"从文本解析测试用例失败: {str(e)}", exc_info=True)
    
    return test_cases

@api_view(['POST'])
@permission_classes([IsAuthenticated])
@authentication_classes([JWTAuthentication])
def export_test_cases(request):
    """
    将测试用例导出为Excel文件
    
    请求参数:
    - test_cases: 测试用例列表
    - filename: 导出文件名，可选
    
    返回:
    - Excel文件下载
    """
    try:
        import pandas as pd
        from django.http import HttpResponse
        from io import BytesIO
        
        # 解析请求数据
        test_cases = request.data.get('test_cases', [])
        filename = request.data.get('filename', 'test_cases.xlsx')
        
        if not test_cases:
            return JsonResponse({
                'code': 400,
                'message': '测试用例不能为空',
                'data': None
            }, status=400)
        
        # 将测试用例转换为DataFrame格式
        cases_data = []
        for case in test_cases:
            case_data = {
                '用例名称': case.get('title', ''),
                '接口路径': case.get('api_path', ''),
                '请求方法': case.get('method', 'GET'),
                '优先级': case.get('priority', '中'),
                '请求头': json.dumps(case.get('headers', {}), ensure_ascii=False),
                '请求体': json.dumps(case.get('body', {}), ensure_ascii=False),
                '断言类型': case.get('assertions', ''),
                '预期结果': json.dumps(case.get('expected_result', {}), ensure_ascii=False),
                '提取器': json.dumps(case.get('extractors', []), ensure_ascii=False)
            }
            cases_data.append(case_data)
        
        # 创建DataFrame
        df = pd.DataFrame(cases_data)
        
        # 创建一个Excel文件
        output = BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='测试用例', index=False)
            
            # 调整列宽
            worksheet = writer.sheets['测试用例']
            for idx, col in enumerate(df.columns):
                max_len = max(df[col].astype(str).map(len).max(), len(col)) + 2
                worksheet.column_dimensions[worksheet.cell(1, idx+1).column_letter].width = max_len
        
        # 设置响应头信息
        response = HttpResponse(output.getvalue(),
                                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
        
    except Exception as e:
        logger.error(f"导出测试用例失败: {str(e)}", exc_info=True)
        return JsonResponse({
            'code': 500,
            'message': f'导出测试用例失败: {str(e)}',
            'data': None
        }, status=500)

@api_view(['POST'])
@permission_classes([IsAuthenticated])
@authentication_classes([JWTAuthentication])
def batch_edit_test_cases(request):
    """
    批量编辑/创建测试用例
    
    请求参数:
    - test_cases: 测试用例列表，每个测试用例包含完整的测试用例数据
    - project_id: 项目ID，如果是创建新用例时使用
    
    返回:
    - 编辑/创建后的测试用例信息列表
    """
    try:
        logger.info(f"接收到批量编辑测试用例请求，用户ID: {request.user.id}")
        
        # 解析请求数据
        test_cases = request.data.get('test_cases', [])
        project_id = request.data.get('project_id', 1)  # 默认使用项目ID 1
        
        if not test_cases:
            logger.warning("请求中未提供测试用例数据")
            return JsonResponse({
                'code': 400,
                'message': '测试用例数据不能为空',
                'data': None
            }, status=400)
        
        logger.debug(f"批量处理{len(test_cases)}个测试用例")
        
        # 导入TestCase模型
        from test_platform.models import TestCase
        
        # 优先级映射
        priority_map = {'高': 2, '中': 1, '低': 0}
        
        # 存储处理结果
        results = []
        created_count = 0
        updated_count = 0
        failed_count = 0
        
        # 批量处理测试用例
        for case_data in test_cases:
            try:
                case_id = case_data.get('id')
                
                # 如果提供了ID，尝试更新现有测试用例
                if case_id:
                    try:
                        test_case = TestCase.objects.get(test_case_id=case_id)
                        logger.debug(f"更新测试用例: ID={case_id}")
                        
                        # 更新测试用例字段
                        if 'title' in case_data:
                            test_case.case_name = case_data.get('title')
                        if 'description' in case_data:
                            test_case.case_description = case_data.get('description', '')
                        if 'api_path' in case_data:
                            test_case.case_path = case_data.get('api_path')
                        if 'method' in case_data:
                            test_case.case_request_method = case_data.get('method')
                        if 'priority' in case_data:
                            priority_value = priority_map.get(case_data.get('priority'), 1)
                            test_case.case_priority = priority_value
                        if 'headers' in case_data:
                            test_case.case_request_headers = json.dumps(case_data.get('headers', {}))
                        if 'body' in case_data:
                            test_case.case_requests_body = json.dumps(case_data.get('body', {}))
                        if 'assertions' in case_data:
                            test_case.case_assert_type = case_data.get('assertions', '')
                        if 'expected_result' in case_data:
                            test_case.case_expect_result = json.dumps(case_data.get('expected_result', {}))
                        if 'extractors' in case_data:
                            test_case.case_extractors = json.dumps(case_data.get('extractors', []))
                        if 'tests' in case_data:
                            test_case.case_tests = json.dumps(case_data.get('tests', []))
                        
                        # 更新时间
                        test_case.update_time = timezone.now()
                        
                        # 保存更新
                        test_case.save()
                        updated_count += 1
                        
                        results.append({
                            'id': test_case.test_case_id,
                            'name': test_case.case_name,
                            'status': 'updated',
                            'success': True
                        })
                    except TestCase.DoesNotExist:
                        logger.warning(f"测试用例不存在: ID={case_id}")
                        failed_count += 1
                        results.append({
                            'id': case_id,
                            'name': case_data.get('title', '未知'),
                            'status': 'failed',
                            'success': False,
                            'message': '测试用例不存在'
                        })
                        continue
                else:
                    # 创建新的测试用例
                    case_name = case_data.get('title', '')
                    logger.debug(f"创建新测试用例: {case_name}")
                    
                    if not case_name:
                        logger.warning("测试用例名称为空，跳过")
                        failed_count += 1
                        results.append({
                            'name': '未命名',
                            'status': 'failed',
                            'success': False,
                            'message': '测试用例名称不能为空'
                        })
                        continue
                    
                    # 映射优先级
                    priority_value = priority_map.get(case_data.get('priority', '中'), 1)
                    
                    # 创建测试用例
                    test_case = TestCase.objects.create(
                        case_name=case_name,
                        case_description=case_data.get('description', ''),
                        case_path=case_data.get('api_path', ''),
                        case_request_method=case_data.get('method', 'GET'),
                        case_priority=priority_value,
                        case_params='',
                        case_status=1,  # 假设1表示"启用"
                        case_precondition='',
                        case_request_headers=json.dumps(case_data.get('headers', {})),
                        case_requests_body=json.dumps(case_data.get('body', {})),
                        case_assert_type=case_data.get('assertions', ''),
                        case_assert_contents=json.dumps(case_data.get('tests', [])),
                        case_expect_result=json.dumps(case_data.get('expected_result', {})),
                        create_time=timezone.now(),
                        update_time=timezone.now(),
                        creator_id=request.user.id,
                        project_id=project_id,
                        case_extractors=json.dumps(case_data.get('extractors', []))
                    )
                    
                    created_count += 1
                    results.append({
                        'id': test_case.test_case_id,
                        'name': test_case.case_name,
                        'status': 'created',
                        'success': True
                    })
            except Exception as e:
                logger.error(f"处理测试用例失败: {str(e)}")
                failed_count += 1
                results.append({
                    'name': case_data.get('title', '未知'),
                    'status': 'failed',
                    'success': False,
                    'message': str(e)
                })
        
        logger.info(f"批量测试用例处理完成: 创建={created_count}, 更新={updated_count}, 失败={failed_count}")
        
        return JsonResponse({
            'code': 200,
            'message': f'测试用例批量处理完成: 创建={created_count}, 更新={updated_count}, 失败={failed_count}',
            'data': {
                'results': results,
                'created_count': created_count,
                'updated_count': updated_count,
                'failed_count': failed_count,
                'total': len(test_cases)
            }
        })
        
    except Exception as e:
        logger.error(f"批量编辑测试用例失败: {str(e)}", exc_info=True)
        return JsonResponse({
            'code': 500,
            'message': f'批量编辑测试用例失败: {str(e)}',
            'data': None
        }, status=500)

@api_view(['POST', 'PUT'])
@permission_classes([IsAuthenticated])
@authentication_classes([JWTAuthentication])
def edit_test_case(request):
    """
    编辑从DeepSeek分析提取的测试用例
    
    请求参数:
    - test_case: 需要编辑的测试用例数据
    - case_id: [可选] 已存在的测试用例ID，如果提供则更新该用例，否则创建新用例
    
    返回:
    - 编辑后的测试用例信息
    """
    try:
        logger.info(f"接收到编辑测试用例请求，用户ID: {request.user.id}")
        
        # 解析请求数据
        test_case_data = request.data.get('test_case', {})
        case_id = request.data.get('case_id')
        
        if not test_case_data:
            logger.warning("请求中未提供测试用例数据")
            return JsonResponse({
                'code': 400,
                'message': '测试用例数据不能为空',
                'data': None
            }, status=400)
        
        # 导入TestCase模型
        from test_platform.models import TestCase
        
        # 如果提供了case_id，则更新现有测试用例
        if case_id:
            logger.info(f"更新现有测试用例，ID: {case_id}")
            try:
                test_case = TestCase.objects.get(test_case_id=case_id)
                
                # 更新测试用例字段
                if 'title' in test_case_data:
                    test_case.case_name = test_case_data.get('title')
                if 'description' in test_case_data:
                    test_case.case_description = test_case_data.get('description', '')
                if 'api_path' in test_case_data:
                    test_case.case_path = test_case_data.get('api_path')
                if 'method' in test_case_data:
                    test_case.case_request_method = test_case_data.get('method')
                if 'priority' in test_case_data:
                    # 映射优先级
                    priority_map = {'高': 2, '中': 1, '低': 0}
                    priority_value = priority_map.get(test_case_data.get('priority'), 1)
                    test_case.case_priority = priority_value
                if 'headers' in test_case_data:
                    test_case.case_request_headers = json.dumps(test_case_data.get('headers', {}))
                if 'body' in test_case_data:
                    test_case.case_requests_body = json.dumps(test_case_data.get('body', {}))
                if 'assertions' in test_case_data:
                    test_case.case_assert_type = test_case_data.get('assertions', '')
                if 'expected_result' in test_case_data:
                    test_case.case_expect_result = json.dumps(test_case_data.get('expected_result', {}))
                if 'extractors' in test_case_data:
                    test_case.case_extractors = json.dumps(test_case_data.get('extractors', []))
                if 'tests' in test_case_data:
                    test_case.case_tests = json.dumps(test_case_data.get('tests', []))
                
                # 更新时间
                test_case.update_time = timezone.now()
                
                # 保存更新
                test_case.save()
                logger.info(f"测试用例更新成功: ID={test_case.test_case_id}, 名称={test_case.case_name}")
                
                return JsonResponse({
                    'code': 200,
                    'message': '测试用例更新成功',
                    'data': {
                        'id': test_case.test_case_id,
                        'name': test_case.case_name,
                        'description': test_case.case_description,
                        'api_path': test_case.case_path,
                        'method': test_case.case_request_method,
                        'priority': {0: '低', 1: '中', 2: '高'}.get(test_case.case_priority, '中'),
                        'headers': json.loads(test_case.case_request_headers) if test_case.case_request_headers else {},
                        'body': json.loads(test_case.case_requests_body) if test_case.case_requests_body else {},
                        'assertions': test_case.case_assert_type,
                        'expected_result': json.loads(test_case.case_expect_result) if test_case.case_expect_result else {},
                        'extractors': json.loads(test_case.case_extractors) if test_case.case_extractors else []
                    }
                })
            except TestCase.DoesNotExist:
                logger.warning(f"测试用例不存在: ID={case_id}")
                return JsonResponse({
                    'code': 404,
                    'message': '测试用例不存在',
                    'data': None
                }, status=404)
        else:
            # 创建新的测试用例
            logger.info("创建新的测试用例")
            
            # 获取必要的字段
            project_id = test_case_data.get('project_id', 1)  # 默认使用项目ID 1
            case_name = test_case_data.get('title', '')
            
            if not case_name:
                logger.warning("测试用例名称为空")
                return JsonResponse({
                    'code': 400,
                    'message': '测试用例名称不能为空',
                    'data': None
                }, status=400)
            
            # 映射优先级
            priority_map = {'高': 2, '中': 1, '低': 0}
            priority_value = priority_map.get(test_case_data.get('priority', '中'), 1)
            
            # 创建测试用例
            test_case = TestCase.objects.create(
                case_name=case_name,
                case_description=test_case_data.get('description', ''),
                case_path=test_case_data.get('api_path', ''),
                case_request_method=test_case_data.get('method', 'GET'),
                case_priority=priority_value,
                case_params='',
                case_status=1,  # 假设1表示"启用"
                case_precondition='',
                case_request_headers=json.dumps(test_case_data.get('headers', {})),
                case_requests_body=json.dumps(test_case_data.get('body', {})),
                case_assert_type=test_case_data.get('assertions', ''),
                case_assert_contents=json.dumps(test_case_data.get('tests', [])),
                case_expect_result=json.dumps(test_case_data.get('expected_result', {})),
                create_time=timezone.now(),
                update_time=timezone.now(),
                creator_id=request.user.id,
                project_id=project_id,
                case_extractors=json.dumps(test_case_data.get('extractors', []))
            )
            
            logger.info(f"测试用例创建成功: ID={test_case.test_case_id}, 名称={test_case.case_name}")
            
            return JsonResponse({
                'code': 200,
                'message': '测试用例创建成功',
                'data': {
                    'id': test_case.test_case_id,
                    'name': test_case.case_name,
                    'description': test_case.case_description,
                    'api_path': test_case.case_path,
                    'method': test_case.case_request_method,
                    'priority': {0: '低', 1: '中', 2: '高'}.get(test_case.case_priority, '中'),
                    'headers': json.loads(test_case.case_request_headers) if test_case.case_request_headers else {},
                    'body': json.loads(test_case.case_requests_body) if test_case.case_requests_body else {},
                    'assertions': test_case.case_assert_type,
                    'expected_result': json.loads(test_case.case_expect_result) if test_case.case_expect_result else {},
                    'extractors': json.loads(test_case.case_extractors) if test_case.case_extractors else []
                }
            })
        
    except Exception as e:
        logger.error(f"编辑测试用例失败: {str(e)}", exc_info=True)
        return JsonResponse({
            'code': 500,
            'message': f'编辑测试用例失败: {str(e)}',
            'data': None
        }, status=500)

@api_view(['POST'])
@permission_classes([IsAuthenticated])
@authentication_classes([JWTAuthentication])
def save_analysis_result(request):
    """
    保存AI分析结果
    
    请求参数:
    - analysis_data: 包含文件名、文件类型、DeepSeek响应、解析的表格数据等
    - project_id: 关联的项目ID (可选)
    
    返回:
    - 保存结果
    """
    try:
        logger.info(f"接收到保存AI分析结果请求，用户ID: {request.user.id}")
        
        # 解析请求数据
        analysis_data = request.data.get('analysis_data', {})
        project_id = request.data.get('project_id')
        
        if not analysis_data:
            logger.warning("请求中未提供分析数据")
            return JsonResponse({
                'code': 400,
                'message': '分析数据不能为空',
                'data': None
            }, status=400)
        
        # 导入AnalysisResult模型
        from test_platform.models import AnalysisResult, Project
        
        # 获取必要字段
        file_name = analysis_data.get('file_name', '')
        file_type = analysis_data.get('file_type', '')
        
        if not file_name or not file_type:
            logger.warning("文件名或文件类型为空")
            return JsonResponse({
                'code': 400,
                'message': '文件名和文件类型不能为空',
                'data': None
            }, status=400)
        
        # 处理DeepSeek响应数据
        deepseek_response = analysis_data.get('deepseek_response')
        sheets = analysis_data.get('sheets', [])
        test_cases = []
        
        # 从DeepSeek响应中提取测试用例
        if deepseek_response and 'test_cases' in deepseek_response:
            test_cases = deepseek_response.get('test_cases', [])
        
        # 创建分析结果记录
        analysis_result = AnalysisResult()
        analysis_result.file_name = file_name
        analysis_result.file_type = file_type
        analysis_result.deepseek_response = json.dumps(deepseek_response) if deepseek_response else '{}'
        analysis_result.sheets_data = json.dumps(sheets) if sheets else '[]'
        analysis_result.test_cases_data = json.dumps(test_cases) if test_cases else '[]'
        analysis_result.creator = request.user
        
        # 关联项目(如果提供)
        if project_id:
            try:
                project = Project.objects.get(project_id=project_id)
                analysis_result.project = project
            except Project.DoesNotExist:
                logger.warning(f"关联的项目不存在: ID={project_id}")
                # 仍然保存分析结果，但不关联项目
        
        # 保存分析结果
        analysis_result.save()
        logger.info(f"AI分析结果保存成功: ID={analysis_result.analysis_id}, 文件名={file_name}")
        
        return JsonResponse({
            'code': 200,
            'message': 'AI分析结果保存成功',
            'data': {
                'analysis_id': analysis_result.analysis_id,
                'file_name': analysis_result.file_name,
                'file_type': analysis_result.file_type,
                'create_time': analysis_result.create_time.strftime('%Y-%m-%d %H:%M:%S'),
                'test_cases_count': len(test_cases)
            }
        })
        
    except Exception as e:
        logger.error(f"保存AI分析结果失败: {str(e)}", exc_info=True)
        return JsonResponse({
            'code': 500,
            'message': f'保存AI分析结果失败: {str(e)}',
            'data': None
        }, status=500)

def process_large_data_with_llm(data, client, model, temperature=0.7, max_chunk_size=8000, system_prompt=None):
    """
    处理大型数据时，将数据分批发送给大模型并汇总结果
    
    :param data: 要处理的数据文本
    :param client: OpenAI客户端实例
    :param model: 使用的模型名称
    :param temperature: 温度参数，控制输出的随机性
    :param max_chunk_size: 每个分块的最大字符数
    :param system_prompt: 系统提示文本
    :return: 汇总后的结果
    """
    logger.info(f"开始处理大型数据，总长度: {len(data)} 字符")
    
    if len(data) <= max_chunk_size:
        logger.info("数据量较小，直接发送处理")
        # 数据量小，直接处理
        system_message = {"role": "system", "content": system_prompt} if system_prompt else {"role": "system", "content": "请分析以下数据。"}
        response = client.chat.completions.create(
            model=model,
            messages=[
                system_message,
                {"role": "user", "content": data}
            ],
            temperature=temperature,
            stream=False
        )
        return response.choices[0].message.content
    
    # 数据量大，分块处理
    logger.info(f"数据量较大，需要分块处理，每块最大 {max_chunk_size} 字符")
    chunks = []
    current_chunk = ""
    
    # 1. 分割数据为合适大小的块
    for line in data.split('\n'):
        if len(current_chunk) + len(line) + 1 > max_chunk_size:
            chunks.append(current_chunk)
            current_chunk = line
        else:
            if current_chunk:
                current_chunk += '\n' + line
            else:
                current_chunk = line
                
    # 添加最后一个块
    if current_chunk:
        chunks.append(current_chunk)
    
    logger.info(f"数据已分割为 {len(chunks)} 个块")
    
    # 2. 逐块处理并收集结果
    partial_results = []
    
    for i, chunk in enumerate(chunks):
        logger.info(f"处理第 {i+1}/{len(chunks)} 块数据，长度: {len(chunk)} 字符")
        
        # 第一块使用完整的系统提示
        if i == 0:
            prompt = system_prompt if system_prompt else "请分析以下数据。这是第一部分数据，之后会有更多部分。请提取关键信息。"
        else:
            prompt = "这是数据的继续部分，请基于之前的分析继续提取信息。"
        
        try:
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": prompt},
                    {"role": "user", "content": chunk}
                ],
                temperature=temperature,
                stream=False
            )
            result = response.choices[0].message.content
            partial_results.append(result)
            logger.info(f"第 {i+1} 块数据处理完成，结果长度: {len(result)} 字符")
        except Exception as e:
            logger.error(f"处理第 {i+1} 块数据时出错: {str(e)}")
            partial_results.append(f"处理出错: {str(e)}")
    
    # 3. 汇总结果
    if len(partial_results) == 1:
        logger.info("只有一个结果块，无需汇总")
        return partial_results[0]
    
    logger.info(f"开始汇总 {len(partial_results)} 个部分结果")
    
    # 将所有部分结果组合为一个字符串
    combined_results = "\n\n===== 部分结果分隔线 =====\n\n".join(partial_results)
    
    # 发送汇总请求
    try:
        summary_response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": """
请对之前分批分析的多个结果进行汇总整合。每个部分结果之间用分隔线隔开。
你需要把所有测试用例整合成一个完整的、一致的JSON数组。
必须严格按照以下格式输出：

```json
[
  {
    "title": "测试用例标题",
    "api_path": "接口路径",
    "method": "请求方法",
    ...其他测试用例字段
  },
  {
    ...第二个测试用例
  }
]
```

请确保输出是有效的JSON格式，只包含测试用例数组，不要添加任何额外的解释。
如果发现重复的测试用例，请只保留一个。
"""},
                {"role": "user", "content": combined_results}
            ],
            temperature=0.2,  # 使用较低的温度确保更确定性的结果
            stream=False
        )
        final_result = summary_response.choices[0].message.content
        logger.info(f"结果汇总完成，最终结果长度: {len(final_result)} 字符")
        return final_result
    except Exception as e:
        logger.error(f"汇总结果时出错: {str(e)}")
        # 如果汇总失败，返回第一个结果
        return partial_results[0]

@api_view(['GET'])
@permission_classes([IsAuthenticated])
@authentication_classes([JWTAuthentication])
def get_analysis_results(request):
    """
    获取AI分析结果
    
    请求参数:
    - analysis_id: [可选] 分析结果ID，如果提供则返回详细结果，否则返回分析结果列表
    - project_id: [可选] 项目ID，如果提供则返回该项目的分析结果列表
    - limit: [可选] 返回结果数量限制，默认为20
    - offset: [可选] 分页偏移量，默认为0
    
    返回:
    - 分析结果列表或详情
    """
    try:
        # 导入模型
        from test_platform.models import AnalysisResult, Project
        
        # 获取请求参数
        analysis_id = request.GET.get('analysis_id')
        project_id = request.GET.get('project_id')
        limit = int(request.GET.get('limit', 20))
        offset = int(request.GET.get('offset', 0))
        
        # 如果提供了analysis_id，则返回详情
        if analysis_id:
            try:
                analysis = AnalysisResult.objects.get(analysis_id=analysis_id)
                
                # 解析JSON字段
                try:
                    deepseek_response = json.loads(analysis.deepseek_response) if analysis.deepseek_response else {}
                except json.JSONDecodeError:
                    deepseek_response = {}
                
                try:
                    sheets_data = json.loads(analysis.sheets_data) if analysis.sheets_data else []
                except json.JSONDecodeError:
                    sheets_data = []
                
                try:
                    test_cases_data = json.loads(analysis.test_cases_data) if analysis.test_cases_data else []
                except json.JSONDecodeError:
                    test_cases_data = []
                
                # 构建响应数据
                data = {
                    'analysis_id': analysis.analysis_id,
                    'file_name': analysis.file_name,
                    'file_type': analysis.file_type,
                    'deepseek_response': deepseek_response,
                    'sheets': sheets_data,
                    'test_cases': test_cases_data,
                    'creator': analysis.creator.username if analysis.creator else None,
                    'project_id': analysis.project.project_id if analysis.project else None,
                    'project_name': analysis.project.name if analysis.project else None,
                    'create_time': analysis.create_time.strftime('%Y-%m-%d %H:%M:%S'),
                    'update_time': analysis.update_time.strftime('%Y-%m-%d %H:%M:%S')
                }
                
                return JsonResponse({
                    'code': 200,
                    'message': '获取分析结果成功',
                    'data': data
                })
            except AnalysisResult.DoesNotExist:
                return JsonResponse({
                    'code': 404,
                    'message': '分析结果不存在',
                    'data': None
                }, status=404)
        
        # 否则，返回分析结果列表
        query = AnalysisResult.objects.all()
        
        # 如果提供了项目ID，则进行过滤
        if project_id:
            query = query.filter(project_id=project_id)
        
        # 获取总记录数
        total = query.count()
        
        # 分页查询
        results = query.order_by('-create_time')[offset:offset+limit]
        
        # 构建响应数据
        result_list = []
        for analysis in results:
            # 提取测试用例数量
            test_cases_count = 0
            if analysis.test_cases_data:
                try:
                    test_cases_data = json.loads(analysis.test_cases_data)
                    test_cases_count = len(test_cases_data) if isinstance(test_cases_data, list) else 0
                except json.JSONDecodeError:
                    pass
                    
            result_list.append({
                'analysis_id': analysis.analysis_id,
                'file_name': analysis.file_name,
                'file_type': analysis.file_type,
                'test_cases_count': test_cases_count,
                'creator': analysis.creator.username if analysis.creator else None,
                'project_id': analysis.project.project_id if analysis.project else None,
                'project_name': analysis.project.name if analysis.project else None,
                'create_time': analysis.create_time.strftime('%Y-%m-%d %H:%M:%S')
            })
        
        return JsonResponse({
            'code': 200,
            'message': '获取分析结果列表成功',
            'data': {
                'total': total,
                'results': result_list
            }
        })
        
    except Exception as e:
        logger.error(f"获取分析结果失败: {str(e)}", exc_info=True)
        return JsonResponse({
            'code': 500,
            'message': f'获取分析结果失败: {str(e)}',
            'data': None
        }, status=500)

@api_view(['GET'])
@permission_classes([IsAuthenticated])
@authentication_classes([JWTAuthentication])
def get_analysis_result_detail(request, analysis_id):
    """
    获取单个AI分析结果详情
    
    :param request: HTTP请求
    :param analysis_id: 分析结果ID
    :return: 详细的分析结果数据
    """
    try:
        # 导入模型
        from test_platform.models import AnalysisResult
        
        try:
            # 查询分析结果
            analysis = AnalysisResult.objects.get(analysis_id=analysis_id)
            
            # 解析JSON字段
            try:
                deepseek_response = json.loads(analysis.deepseek_response) if analysis.deepseek_response else {}
            except json.JSONDecodeError:
                deepseek_response = {}
            
            try:
                sheets_data = json.loads(analysis.sheets_data) if analysis.sheets_data else []
            except json.JSONDecodeError:
                sheets_data = []
            
            try:
                test_cases_data = json.loads(analysis.test_cases_data) if analysis.test_cases_data else []
            except json.JSONDecodeError:
                test_cases_data = []
            
            # 确保deepseek_response包含test_cases
            if 'test_cases' not in deepseek_response:
                deepseek_response['test_cases'] = test_cases_data
            
            # 构建响应数据，按照前端要求的格式
            data = {
                'id': analysis.analysis_id,
                'file_name': analysis.file_name,
                'file_type': analysis.file_type,
                'model': deepseek_response.get('model', 'deepseek-chat'),
                'create_time': analysis.create_time.isoformat(),
                'deepseek_response': deepseek_response,
                'sheets': sheets_data
            }
            
            return JsonResponse({
                'code': 200,
                'message': 'success',
                'data': data
            })
        except AnalysisResult.DoesNotExist:
            return JsonResponse({
                'code': 404,
                'message': '分析结果不存在',
                'data': None
            }, status=404)
        
    except Exception as e:
        logger.error(f"获取分析结果详情失败: {str(e)}", exc_info=True)
        return JsonResponse({
            'code': 500,
            'message': f'获取分析结果详情失败: {str(e)}',
            'data': None
        }, status=500)
